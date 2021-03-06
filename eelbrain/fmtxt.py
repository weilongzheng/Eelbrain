# -*- coding: utf-8 -*-
# Author:  Christian Brodbeck <christianbrodbeck@nyu.edu>
"""Document model for formatted text documents

Objects for the abstract representation of a document which can be into
different formats. Currently (incomplete) support for str, RTF, TeX and HTML.

Base classes:

:class:`FMTextElement`
    A string (or object with a str representation) along with formatting
    information.
:class:`FMText`
    An FMTextElement whose contents is a list of FMText objects.


Subclasses for specific purposes:

:class:`Table`
    Tables with multicolumn cells.
:class:`Image`
    Images.
:class:`Figure`
    Figure with content and caption.
:class:`Section`
    Document section containing a title and content (any other FMText objects,
    including other Section objects for subsections).
:class:`Report`
    Document consisting of several sections plus a title.

Whenever an parameter asks for an FMText object, :func:`asfmtxt` handles the
coercion into an appropriate FMTextElement subclass.

FMText objects provide an interface to formatting through different methods:

- the :py:meth:`__str__` method for a string representation
- a :py:meth:`get_tex` method for a TeX representation
- a :py:meth:`get_html` method for a HTML representation

The module also provides functions that work with fmtxt objects:

- :func:`save_tex` for saving an object's tex representation
- :func:`copy_tex` for copying an object's tex representation to
  the clipboard
- :func:`save_pdf` for saving a pdf
- :func:`copy_pdf` for copying a pdf to the clipboard
- :func:`save_html` for saving an HTML file

"""
import base64
import datetime
from html.parser import HTMLParser
from importlib import import_module
from itertools import repeat
from math import ceil
import os
import pickle
import re
import shutil
import socket
from io import BytesIO
import tempfile
import time
from typing import Union, Iterable, List as ListType, Tuple
import webbrowser

import numpy as np
import matplotlib.figure
from matplotlib.image import imsave
from matplotlib.mathtext import math_to_image

from ._utils.tex import latex2pdf
from ._utils import ui


# types
FMTextLike = Union['FMTextElement', str, ListType['FMTextLike'], Tuple['FMTextLike']]


preferences = dict(
                   keep_recent=3,  # number of recent tables to keep in memory
                   html_tables_in_fig=True,
                   )

# Tags
# ----
# Formatting instructions, actions depend on output format.
# Two ways of dealing with them:
#  - substitute tag with X_TAGS[tag]
#  - format native format string with content (X_SUBS[tag] % content)
#
# str:
#     ``sub % content`` if the tag is in _STR_SUBS, otherwise it is
#     ignored
# TeX:
#     Ignor a tag if it is in _TEX_IGNORE;
#     ``sub % content`` if the tag is in _TEX_SUBS,
#     otherwise ``"%s{%s}" % (tag, content)``
# HTML:
#     Formatted HTML if the tag is in _HTML_TAGS, otherwise it is ignored
# RTF:
#     Formatted RTF if the tag is in _RTF_SUBS, otherwise it is ignored

_STR_SUBS = {r'_': '(%s)'}

_TEX_IGNORE = ('font',)
_TEX_SUBS = {'paragraph': "\n\n%s\n\n",
             'math': "$%s$"}  # LaTeX math but normal HTML

_HTML_TAGS = {r'_': 'sub',
              r'^': 'sup',
              r'\author': 'author',
              r'\emph': 'em',
              r'\textbf': 'b',
              r'\textit': 'i',
              'paragraph': 'p',
              'code': 'code',
              'font': 'font'}

_RTF_SUBS = {r'\emph': "\i %s\i0"}

_html_alignments = {'l': 'left',
                    'r': 'right',
                    'c': 'center'}

_html_reserved_chars = {'<': '&lt;', '>': '&gt;', '&': '&amp;'}
_html_escape_pattern = re.compile('|'.join(map(re.escape, _html_reserved_chars)))

_tex_escape_chars = {
    '&': r'\&',
    '%': r'\%',
    '$': r'\$',
    '#': r'\#',
    '_': r'\_',
    '{': r'\{',
    '}': r'\}',
    '[': '{[}',
    ']': '{]}',
    '"':  "{''}",
    '~': r'\textasciitilde{}',
    '^': r'\^{}',
    '\\': r'\textbackslash{}',
    '<': r'\textless{}',
    '>': r'\textgreater{}',
}
_tex_escape_pattern = re.compile('|'.join(map(re.escape, _tex_escape_chars)))


def _html_repl(m):
    return _html_reserved_chars[m.group(0)]


def escape_html(text):
    return _html_escape_pattern.sub(_html_repl, text)#.encode('ascii', 'xmlcharrefreplace')


def _tex_repl(m):
    return _tex_escape_chars[m.group(0)]


def escape_tex(text):
    return _tex_escape_pattern.sub(_tex_repl, text)


STYLE = """
.float {
    float:left
}
"""

_html_doc_template = """<!DOCTYPE html>
<html>
<head>
    {meta}<title>{title}</title>{style}
</head>

<body>

{body}

</body>
</html>
"""

# to keep track of recent tex out and allow copying
_recent_texout = []


def _add_to_recent(tex_obj):
    keep_recent = preferences['keep_recent']
    if keep_recent:
        if len(_recent_texout) >= keep_recent - 1:
            _recent_texout.pop(0)
        _recent_texout.append(tex_obj)


def get_pdf(tex_obj):
    "Generate PDF from an FMText object (using :mod:`tex`)"
    if isinstance(tex_obj, (Report, Section)):
        doc_class = '{article}'
        standalone = False
    else:
        doc_class = '[border=2pt]{standalone}'
        standalone = True
    txt = tex(tex_obj, {'standalone': standalone})
    document = """
\\documentclass%s
\\usepackage{booktabs}
\\begin{document}
%s
\\end{document}
""" % (doc_class, txt)
    pdf = latex2pdf(document)
    return pdf


def rtf_document(fmtext):
    return "{\\rtf1\\ansi\\deff0\n\n%s\n}" % fmtext.get_rtf()


def save_html(fmtxt, path=None, embed_images=True, meta=None):
    """Save an FMText object in HTML format

    Parameters
    ----------
    fmtext : FMText
        Object to save.
    path : str (optional)
        Destination filename. If unspecified, a file dialog will open to ask
        for a destination.
    embed_images : bool
        Embed images in the HTML file (default True). If False, a separate
        folder containing image files is created.
    meta : dict
        Meta-information for document head.
    """
    if path is None:
        msg = "Save as HTML"
        path = ui.ask_saveas(msg, msg, [('HTML (*.html)', '*.html')])
        if not path:
            return
    path = os.path.abspath(path)

    extension = '.html'
    if path.endswith(extension):
        resource_dir = path[:-len(extension)]
        file_path = path
    else:
        resource_dir = path
        file_path = path + extension
    root = os.path.dirname(file_path)

    if embed_images:
        resource_dir = None
    else:
        if os.path.exists(resource_dir):
            shutil.rmtree(resource_dir)
        os.mkdir(resource_dir)
        resource_dir = os.path.relpath(resource_dir, root)

    buf = make_html_doc(fmtxt, root, resource_dir, meta=meta)
    buf_enc = buf.encode('ascii', 'xmlcharrefreplace')
    with open(file_path, 'wb') as fid:
        fid.write(buf_enc)


def save_pdf(tex_obj, path=None):
    "Save an fmtxt object as a pdf"
    pdf = get_pdf(tex_obj)
    if path is None:
        msg = "Save as PDF"
        path = ui.ask_saveas(msg, msg, [('PDF (*.pdf)', '*.pdf')])
    if path:
        with open(path, 'wb') as f:
            f.write(pdf)


def save_rtf(fmtext, path=None):
    """Save an FMText object in Rich Text format

    Parameters
    ----------
    fmtext : FMText
        Object to save.
    path : str (optional)
        Destination filename. If unspecified, a file dialog will open to ask
        for a destination.
    """
    text = rtf_document(fmtext)
    if path is None:
        path = ui.ask_saveas("Save RTF", filetypes=[('Rich Text File (*.rtf)', '*.rtf')])
    if path:
        with open(path, 'w') as fid:
            fid.write(text)


def save_tex(tex_obj, path=None):
    "Save an FMText object as a PDF"
    txt = tex(tex_obj)
    if path is None:
        path = ui.ask_saveas("Save tex", filetypes=[('tex', 'tex source code')])
    if path:
        with open(path, 'w') as f:
            f.write(txt)


def _save_txt(text, path=None):
    if path is None:
        path = ui.ask_saveas("Save Text File",
                             filetypes=[("Plain Text File (*.txt)", "*.txt")])
    if path:
        with open(path, 'w') as fid:
            fid.write(text)


def copy_pdf(tex_obj=-1):
    """Copt an fmtxt object to the clipboard as PDF.

    Parameters
    ----------
    tex_obj : FMText | int
        Either an FMText object that can be rendered, or an ``int`` to retrieve
        an FMText item from a list of recently displayed FMText objects.
    """
    if isinstance(tex_obj, int):
        tex_obj = _recent_texout[tex_obj]

    # save pdf to temp file
    pdf = get_pdf(tex_obj)
    fd, path = tempfile.mkstemp('.pdf')
    os.write(fd, pdf)
    os.close(fd)

    # copy to clip-board
    ui.copy_file(path)


def copy_tex(tex_obj):
    "Copy an fmtxt object to the clipboard as tex code"
    txt = tex(tex_obj)
    ui.copy_text(txt)


def html(text, env={}):
    """Generate HTML for any object with a string representation

    Parameters
    ----------
    text : any
        Object to be converted to HTML. If the object has a ``.get_html()``
        method the result of this method is returned, otherwise ``str(text)``.
    env : dict
        Environment for HTML.
    """
    if hasattr(text, 'get_html'):
        return text.get_html(env)
    else:
        return str(text)


def make_html_doc(body, root, resource_dir=None, title=None, meta=None):
    """Generate HTML document

    Parameters
    ----------
    body : fmtxt-object
        FMTXT object which should be formatted into an HTML document.
    root : str
        Path to the directory in which the HTML file is going to be located.
    resource_dir : None | str
        Path to the directory containing resources like images, relative to
        root. If None, images are embedded.
    title : None | FMText
        Document title. The default is to try to infer the title from the body
        or use "Untitled".
    meta : dict
        Meta-information for document head.

    Returns
    -------
    html : str
        HTML document.
    """
    if title is None:
        if hasattr(body, '_site_title') and body._site_title is not None:
            title = html(body._site_title)
        elif hasattr(body, '_heading'):
            title = html(body._heading)
        else:
            title = "Untitled"

    if meta:
        meta = '<meta %s>\n' % ' '.join('%s=%r' % x for x in meta.items())
    else:
        meta = ''

    style = '\n'.join(('', '<style>', STYLE, '</style>'))

    env = {'root': root, 'resource_dir': resource_dir}
    txt_body = html(body, env)
    return _html_doc_template.format(meta=meta, title=title, style=style,
                                     body=txt_body)


def tex(text, env=None):
    """Create html code for any object with a string representation

    Parameters
    ----------
    text : any
        Object to be converted to HTML. If the object has a ``.get_html()``
        method the result of this method is returned, otherwise ``str(text)``.
    env : dict
        Environment for FMTXT compilation.
    """
    if hasattr(text, 'get_tex'):
        if env is None:
            env = {'math': False}
        elif 'math' not in env:
            env['math'] = False
        return text.get_tex(env)
    else:
        return str(text)


_html_temp = '<{tag}>{body}</{tag}>'
_html_temp_opt = '<{tag} {options}>{body}</{tag}>'


def _html_element(tag, body, env, options=None):
    """Format an HTML element

    Parameters
    ----------
    tag : str
        The HTML tag.
    body : FMText
        The main content between the tags.
    env : dict
        Environment for FMTXT compilation.
    options : dict
        HTML options to be inserted in the start tag.
    """
    if options:
        opt = ' '.join('%s="%s"' % (k, options[k]) for k in sorted(options))
        txt = _html_temp_opt.format(tag=tag, options=opt, body=html(body, env))
    else:
        txt = _html_temp.format(tag=tag, body=html(body, env))
    return txt


def asfmtext(content, tag=None):
    "Convert non-FMText objects to FMText"
    if isinstance(content, (FMTextElement, FMTextConstant)):
        if tag:
            return FMTextElement(content, tag)
        else:
            return content
    elif hasattr(content, '_asfmtext'):
        return asfmtext(content._asfmtext(), tag)
    elif isinstance(content, (list, tuple)):
        return FMText(content, tag)
    elif isinstance(content, matplotlib.figure.Figure):
        image = Image()
        content.savefig(image)
        return image
    else:
        return Text(content, tag)


def asfmtext_or_none(content, tag=None):
    if content is None:
        return None
    else:
        return asfmtext(content, tag)


class FMTextConstant(object):

    def __init__(self, name, html, rtf, tex, text):
        self.name = name
        self.html = html
        self.rtf = rtf
        self.tex = tex
        self.text = text

    def __repr__(self):
        return self.name

    def get_html(self, env):
        return self.html

    def get_rtf(self, env):
        return self.rtf

    def get_tex(self, env):
        return self.tex

    def get_str(self, env):
        return self.text


linebreak = FMTextConstant(
    name='linebreak',
    html='<br style="clear:left">\n',
    rtf='\\line\n',
    tex='\\\\\n',
    text='\n',
)


class FMTextElement(object):
    """Represent a text element along with formatting properties.

    The elementary unit of the :py:mod:`fmtxt` module. It can function as a
    string, but can hold formatting properties such as font properties.

    The following methods are used to get different string representations:

     - FMText.get_str() -> unicode
     - FMText.get_tex() -> str (TeX)
     - FMText.get_html() -> str (HTML)
     - str(FMText) -> str

    """
    def __init__(self, content, tag=None, options=None):
        """Represent a value along with formatting properties.

        Parameters
        ----------
        content : object
            Any item with a string representation (str, scalar, ...).
        tag : str
            Formatting tag.
        options : dict
            Options for HTML tags.
        """
        self.content = content
        self.tag = tag
        self.options = options

    def __repr__(self):
        name = self.__class__.__name__
        args = ', '.join(self._repr_items())
        return "%s(%s)" % (name, args)

    def _repr_items(self):
        items = [repr(self.content)]
        if self.tag:
            items.append(repr(self.tag))
            if self.options:
                items.append(repr(self.options))
        return items

    def __str__(self):
        return self.get_str()

    def __add__(self, other):
        if isinstance(other, str) and other == '':
            # added to prevent matplotlib from thinking Image is a file path
            raise ValueError("Can't add empty string")

        return FMText([self, other])

    def copy_pdf(self):
        "Copy PDF to clipboard"
        copy_pdf(self)

    def copy_tex(self):
        "Copy TeX to clipboard"
        copy_tex(self)

    def _get_core(self, env):
        "Unicode core"
        return str(self.content)

    def get_html(self, env):
        "Complete HTML representation"
        txt = self._get_html_core(env)

        if self.tag and self.tag in _HTML_TAGS:
            tag = _HTML_TAGS[self.tag]
            return _html_element(tag, txt, env, self.options)

        return txt

    def _get_html_core(self, env):
        "HTML representation of everything inside the tag"
        return escape_html(self._get_core(env))

    def get_rtf(self, env={}):
        if self.tag in _RTF_SUBS:
            return _RTF_SUBS[self.tag] % self._get_rtf_core(env)
        else:
            return self._get_rtf_core(env)

    def _get_rtf_core(self, env):
        return self._get_core(env)

    def get_str(self, env={}):
        "Unicode representation"
        if self.tag in _STR_SUBS:
            return _STR_SUBS[self.tag] % self._get_str_core(env)
        else:
            return self._get_str_core(env)

    def _get_str_core(self, env):
        return self._get_core(env)

    def get_tex(self, env):
        if self.tag == 'math':
            if env['math']:
                raise RuntimeError("Nested math tag")
            env['math'] = True
        txt = self._get_tex_core(env)
        if self.tag == 'math':
            env['math'] = False

        if self.tag and self.tag not in _TEX_IGNORE:
            if self.tag in _TEX_SUBS:
                txt = _TEX_SUBS[self.tag] % txt
            else:
                txt = r"%s{%s}" % (self.tag, txt)

        return txt

    def _get_tex_core(self, env):
        out = self._get_core(env)
        if env['math']:
            return out
        else:
            return escape_tex(out)

    def save_html(self, path=None, embed_images=True, meta=None):
        """Save in HTML format

        Parameters
        ----------
        path : str (optional)
            Destination filename. If unspecified, a file dialog will open to ask
            for a destination.
        embed_images : bool
            Embed images in the HTML file (default True). If False, a separate
            folder containing image files is created.
        meta : dict
            Meta-information for document head.
        """
        save_html(self, path, embed_images, meta)

    def save_pdf(self, path=None):
        """Save in PDF format

        Parameters
        ----------
        path : str (optional)
            Destination filename. If unspecified, a file dialog will open to ask
            for a destination.
        """
        save_pdf(self, path)

    def save_rtf(self, path=None):
        """Save in Rich Text format

        Parameters
        ----------
        path : str (optional)
            Destination filename. If unspecified, a file dialog will open to ask
            for a destination.
        """
        save_rtf(self, path)

    def save_tex(self, path=None):
        """Save in TeX format

        Parameters
        ----------
        path : str (optional)
            Destination filename. If unspecified, a file dialog will open to ask
            for a destination.
        """
        save_tex(self, path)

    def save_txt(self, path=None):
        """Save as plain text file

        Parameters
        ----------
        path : str (optional)
            Destination filename. If unspecified, a file dialog will open to ask
            for a destination.
        """
        _save_txt(self.get_str(), path)


class FMText(FMTextElement):
    r"""List of :class:`FMTextElement` items

    Parameters
    ----------
    content : FMTextLike
        Any item with a string representation (str, FMText, scalar, ...)
        or an object that iterates over such items (e.g. a list of FMText).
    tag : str
        Formatting tag.
    options : dict
        Options for HTML tags.

    Notes
    -----
    The ``property`` argument is primarily used for TeX commands; for
    example, ``txt = FMText('Brie', r'\textbf')`` will result in the
    following TeX representation: r"\textbf{Brie}". Some properties are
    automatically translated into HTML; thus, the HTML representation of
    ``txt`` would be "<b>Brie</b>". The following are additional proprties
    that don't exist in TeX:

    'paragraph'
        A <p> tag in HTML, and simple line breaks in TeX.
    """
    def __init__(
            self,
            content: FMTextLike = None,
            tag: str = None,
            options: dict = None,
    ):
        if content is None:
            content = []
        elif isinstance(content, (list, tuple)):
            content = [asfmtext(item) for item in content]
        else:
            content = [asfmtext(content)]
        FMTextElement.__init__(self, content, tag, options)

    def append(self, content):
        """Append content to the FMText item

        Parameters
        ----------
        content : str | object | iterable
            Any item with a string representation (str, FMText, scalar, ...)
            or an object that iterates over such items (e.g. a list of FMText).
        """
        self.content.append(asfmtext(content))

    def _get_html_core(self, env):
        return ''.join(i.get_html(env) for i in self.content)

    def _get_rtf_core(self, env):
        return ''.join(i.get_rtf(env) for i in self.content)

    def _get_str_core(self, env):
        return ''.join(i.get_str(env) for i in self.content)

    def _get_tex_core(self, env):
        return ''.join(i.get_tex(env) for i in self.content)


class Code(FMTextElement):
    """Code (block) element

    Parameters
    ----------
    code : str
        Multiline string to be displayed as code.
    """
    def __init__(self, content):
        assert isinstance(content, str)
        FMTextElement.__init__(self, content, 'code')

    def get_tex(self, env):
        raise NotImplementedError

    def _get_html_core(self, env):
        return linebreak.html.join(line.rstrip().replace(' ', '&nbsp;') for
                                   line in self.content.splitlines())


class Text(FMTextElement):

    def __init__(self, content, tag=None):
        if not isinstance(content, str):
            content = str(content)
        FMTextElement.__init__(self, content, tag)


class Link(FMTextElement):

    def __init__(self, content: FMTextLike, url: str):
        FMTextElement.__init__(self, content)
        self.url = str(url)

    def get_html(self, env):
        return '<a href="%s">%s</a>' % (self.url, FMTextElement.get_html(self, env))


class Number(FMTextElement):

    def __init__(self, content, tag=None, fmt='%s', drop0=False):
        if not np.isscalar(content):
            if getattr(content, 'ndim', None) == 0:
                content = content.item()
            else:
                raise TypeError("content=%s" % repr(content))
        FMTextElement.__init__(self, content, tag)
        self.fmt = fmt
        self.drop0 = drop0

    def _repr_items(self):
        items = FMTextElement._repr_items(self)
        if self.fmt != '%s':
            items.append('fmt=%r' % self.fmt)
        if self.drop0:
            items.append('drop0=True')
        return items

    def _get_core(self, env):
        if np.isnan(self.content):
            return 'NaN'
        elif isinstance(self.content, (bool, np.bool_, np.bool8)):
            return str(self.content)
        else:
            fmt = env.get('fmt', self.fmt)
            txt = fmt % self.content
            if self.drop0:
                if txt.startswith('0.'):
                    return txt[1:]
                elif txt.startswith('-0.'):
                    return '-' + txt[2:]
            return txt


class P(Number):

    def __init__(self, content):
        Number.__init__(self, content, fmt='%.3f', drop0=True)

    def _get_core(self, env):
        if self.content < .001:
            return '< .001'
        else:
            return Number._get_core(self, env)


class Math(FMTextElement):
    r""":class:`FMTextElement` for math expressions

    Parameters
    ----------
    content : str
        LaTeX math expression (excluding '$'). E.g. ``r'\sqrt{a}'``.
    equation : bool
        Whether to display the expression as a separate equation (as
        oppposed to inline).
    """
    def __init__(self, content, equation=False):
        FMTextElement.__init__(self, content)
        self._equation = equation

    def __repr__(self):
        items = [repr(self.content)]
        if self._equation:
            items.append("equation=True")
        return "Math(%s)" % ', '.join(items)

    def get_html(self, env={}):
        im = Image("LaTeX Equation", 'svg', self.content)
        math_to_image("$%s$" % self.content, im, format='svg')
        if self._equation:
            return '\n<br>\n%s\n<br>\n' % im.get_html(env)
        else:
            return im.get_html(env)

    def get_tex(self, env):
        if self._equation:
            return "\\begin{equation}$%s$\\end{equation}" % self.content
        else:
            return "$%s$" % self.content


class EquationArray(FMTextElement):
    """:class:`FMTextElement` for equation arrays

    Parameters
    ----------
    eqnarray : tuple of str
        Tuple of lines for the equation array.
    """
    def __init__(self, eqnarray):
        FMTextElement.__init__(self, eqnarray)

    def __repr__(self):
        return "EquationArray(%s)" % repr(self.content)

    def get_html(self, env):
        ims = []
        for line in self.content:
            im = Image("LaTeX EquationArray", 'svg', line)
            math_to_image("$%s$" % line.replace('&', ''), im, format='svg')
            ims.append(im.get_html(env))
        return '\n<br>\n%s\n<br>\n' % '\n<br>\n'.join(ims)

    def get_tex(self, env):
        return "\\begin{eqnarray}%s\\end{eqnarray}" % r'\\'.join(self.content)


class Stars(FMTextElement):
    """:class:`FMTextElement` for decoration of p-values

    Shortcut for adding stars to a table and spaces in place of absent stars,
    so that alignment to the right can be used.
    """
    def __init__(self, n, of=3, tag="^"):
        if isinstance(n, str):
            self.n = len(n.strip())
        else:
            self.n = n
        self.of = of
        if np.isreal(n):
            text = '*' * n + ' ' * (of - n)
        else:
            text = n.ljust(of)
        FMTextElement.__init__(self, text, tag)

    @classmethod
    def from_p(cls, p, of=3, tag='^'):
        n = sum((p <= 0.001, p <= 0.01, p <= 0.05))
        return cls(n, of, tag)

    def _get_tex_core(self, env):
        txt = self._get_core(env)
        spaces = r'\ ' * (self.of - self.n)
        return txt + spaces


class List(FMTextElement):
    """Bulletted list of FMText elements"""
    def __init__(
            self,
            head: FMTextLike = None,
            items: Iterable[FMTextLike] = None,
            ordered: bool = False,
    ):
        """Bulletted list of FMText elements

        Parameters
        ----------
        head : FMTextLike
            First line on higher level (no bullet for highest list, or list
            element for subordinate list).
        items : iterable of FMTextLike
            List items.
        ordered : bool
            Whether to use the "ol" HTML tag (instead of "ul").
        """
        self.ordered = ordered
        self.head = asfmtext_or_none(head)
        if items is None:
            self.items = []
        else:
            self.items = list(map(asfmtext, items))

    def _repr_items(self):
        if self.ordered:
            return [repr(self.head), repr(self.items), repr(self.ordered)]
        elif self.items:
            return [repr(self.head), repr(self.items)]
        elif self.head:
            return [repr(self.head)]
        else:
            return []

    def add_item(self, item: FMTextLike):
        "Add an item to the list"
        self.items.append(asfmtext(item))

    def add_sublist(
            self,
            head: FMTextLike = None,
            items: Iterable[FMTextLike] = None,
            ordered: bool = False,
    ):
        """Add an item with a subordinate list

        Parameters
        ----------
        head : FMTextLike
            Text for the parent item
        items : iterable of FMTextLike
            Subordinate list items.
        ordered : None | bool
            Whether to use the "ol" HTML tag (instead of "ul"). The default is
            to inherit the parent list's setting.

        Returns
        -------
        sublist : List
            The subordinate list.
        """
        if ordered is None:
            ordered = self.ordered
        sublist = List(head, items, ordered)
        self.add_item(sublist)
        return sublist

    def get_html(self, env={}):
        items = []
        if self.head is not None:
            items.append(self.head.get_html(env))
        tag = 'ol' if self.ordered else 'ul'
        items.append('<%s>' % tag)

        # body
        for item in self.items:
            items.append(_html_element('li', item, env))

        items.append('</%s>' % tag)
        return '\n'.join(items)

    def get_str(self, env={}):
        out = []
        if self.head is not None:
            out.append(self.head.get_str(env))

        for item in self.items:
            if isinstance(item, List):
                lines = item.get_str(env).splitlines()
                out.append('- %s' % lines[0])
                out.extend('  %s' % line for line in lines[1:])
            else:
                out.append('- %s' % str(item))
        return '\n'.join(out)


# Table ---

class Cell(FMText):

    def __init__(self, content='', tag=None, width=1, just=None):
        """A cell for a table

        Parameters
        ----------
        text : FMText
            Cell content.
        width : int
            Width in columns for multicolumn cells.
        just : None | 'l' | 'r' | 'c'
            Justification. None: use column standard.
        others :
            FMText parameters.
        """
        FMText.__init__(self, content, tag)
        self.width = width
        if width > 1 and not just:
            self.just = 'l'
        else:
            self.just = just

    @classmethod
    def coerce(cls, content):
        if isinstance(content, cls):
            return content
        return cls(content)

    def _repr_items(self):
        items = FMText._repr_items(self)
        if self.width != 1:
            i = min(2, len(items))
            items.insert(i, 'width=%s' % self.width)
        return items

    def __len__(self):
        return self.width

    def get_html(self, env={}):
        html_repr = FMText.get_html(self, env)
        options = []
        if self.width > 1:
            options.append('colspan="%i"' % self.width)
        if self.just:
            align = _html_alignments[self.just]
            options.append('align="%s"' % align)

        if options:
            start_tag = '<td %s>' % ' '.join(options)
        else:
            start_tag = '<td>'

        html_repr = ' %s%s</td>' % (start_tag, html_repr)
        return html_repr

    def get_rtf(self, env={}):
        return "%s\\intbl\\cell" % FMText.get_rtf(self, env)

    def get_tex(self, env={}):
        tex_repr = FMText.get_tex(self, env)
        if self.width > 1 or self.just:
            tex_repr = r"\multicolumn{%s}{%s}{%s}" % (self.width, self.just,
                                                      tex_repr)
        return tex_repr


class Row(list):
    """Row for a Table"""
    def __init__(self, n_columns, items=()):
        self.n_columns = n_columns
        list.__init__(self, (Cell.coerce(item) for item in items))

    @classmethod
    def coerce(cls, obj, n_columns):
        if isinstance(obj, cls):
            if obj.n_columns != n_columns:
                raise ValueError("Required Row with %i columns, got %i" %
                                 (n_columns, obj.n_columns))
            return obj
        return cls(n_columns, obj)

    def __len__(self):
        return sum([len(cell) for cell in self])

    def __setitem__(self, key, value):
        if isinstance(key, slice):
            start, stop, stride = key.indices(self.n_columns)
            for i, v in zip(range(start, stop, stride), value):
                self[i] = v
            return
        elif not isinstance(key, int):
            raise TypeError("Row index %r" % (key,))

        column = key if key >= 0 else key + self.n_columns
        if not 0 <= column < self.n_columns:
            raise IndexError("%i for row with %i columns" % (key, self.n_columns))

        cell = Cell.coerce(value)
        if len(self) <= column:
            while len(self) < column:
                self.append(Cell(''))
            self.append(cell)
            return

        column_i = 0
        for i, cell_i in enumerate(self):
            if column_i == column:
                list.__setitem__(self, i, cell)
                return
            column_i += len(cell_i)
        raise IndexError("Column %i is part of a multi-column cell" % (column,))

    def __repr__(self):
        return "Row(%s)" % list.__repr__(self)

    def __str__(self):
        return ' '.join([str(cell) for cell in self])

    def _strlen(self, env):
        "List of cell-str-lengths; multicolumns are handled poorly"
        lens = []
        for cell in self:
            cell_len = len(cell.get_str(env))
            n_columns = len(cell)
            if n_columns == 1:
                lens.append(cell_len)
            else:
                # TODO: better handling of multicolumn
                col_len = int(ceil(cell_len / n_columns))
                lens.extend(repeat(col_len, n_columns))
        return lens

    def cell(
            self,
            content: FMTextLike = None,
            tag: str = None,
            width: int = 1,
            just: str = None,
    ):
        """Add a cell to the row

        Parameters
        ----------
        content : FMText
            Cell content.
        tag : str
            Formatting tag.
        width : int
            Width in columns for multicolumn cells.
        just : 'l' | 'r' | 'c'
            Justification (default: use column standard).
        """
        cell = Cell(content, tag, width, just)
        if len(cell) + len(self) > self.n_columns:
            raise ValueError(f"width={cell.width}: exceeds table width")
        self.append(cell)

    def cells(self, *cells, **kwargs):
        "Add several simple cells with one command"
        for cell in cells:
            self.cell(cell, **kwargs)

    def get_html(self, env={}):
        html = '\n'.join(cell.get_html(env) for cell in self)
        html = '<tr>\n%s\n</tr>' % html
        return html

    def get_rtf(self, env={}):
        return '\n'.join([cell.get_rtf(env) for cell in self] + ['\\row'])

    def get_str(self, c_width, c_just, delimiter='   ', env={}):
        "String of the row using column spacing ``c_width``"
        col = 0
        out = []
        for cell in self:
            if cell.width == 1:
                strlen = c_width[col]
                if cell.just:
                    just = cell.just
                else:
                    just = c_just[col]
            else:
                strlen = sum(c_width[col:col + cell.width])
                strlen += len(delimiter) * (cell.width - 1)
                just = cell.just
            col += cell.width
            txt = cell.get_str(env)
            if just == 'l':
                txt = txt.ljust(strlen)
            elif just == 'r':
                txt = txt.rjust(strlen)
            elif just == 'c':
                rj = strlen // 2
                txt = txt.rjust(rj).ljust(strlen)
            out.append(txt)
        return delimiter.join(out)

    def get_tex(self, env=None):
        if env is None:
            return tex(self)
        out = ' & '.join(cell.get_tex(env) for cell in self)
        out += r" \\"
        return out

    def get_tsv(self, delimiter, fmt=None):
        env = {'fmt': fmt}
        txt = delimiter.join(cell.get_str(env) for cell in self)
        return txt


class Table(FMTextElement):
    r"""A table :class:`FMText` element

    Parameters
    ----------
    columns : str
        alignment for each column, e.g. ``'lrr'``
    rules : bool
        Add toprule and bottomrule
    title : None | text
        Title for the table.
    caption : None | text
        Caption for the table.
    rows : list of Row
        Table body.

    Examples
    --------
    >>> from eelbrain import fmtxt
    >>> table = fmtxt.Table('lll')
    >>> table.cell()
    >>> table.cell("example 1")
    >>> table.cell("example 2")
    >>> table.midrule()
    >>> table.cell("string")
    >>> table.cell('???')
    >>> table.cell('another string')
    >>> table.cell("Number")
    >>> table.cell(4.5)
    >>> table.cell(2./3, fmt='%.4g')
    >>> print(table)
             example 1   example 2
    -----------------------------------
    string   ???         another string
    Number   4.5         0.6667
    >>> print(table.get_tex())
    \begin{center}
    \begin{tabular}{lll}
    \toprule
     & example 1 & example 2 \\
    \midrule
    string & ??? & another string \\
    Number & 4.5 & 0.6667 \\
    \bottomrule
    \end{tabular}
    \end{center}
    >>> table.save_tex()

    """
    def __init__(self, columns, rules=True, title=None, caption=None, rows=[]):
        self.columns = columns
        self.n_columns = len(columns)
        self.rows = rows[:]
        self.rules = rules
        self.title(title)
        self.caption(caption)
        self._active_row = None

    @property
    def shape(self):
        return len(self.rows), self.n_columns

    def __len__(self):
        return len(self.rows)

    def __getitem__(self, item):
        if isinstance(item, int):
            return self.rows[item]
        elif isinstance(item, slice):
            rows = self.rows[item]
            return Table(self.columns, rules=self.rules, title=self._title,
                         caption=self._caption, rows=rows)
        raise TypeError(f'Table index {item}')

    def __setitem__(self, key, value):
        if isinstance(key, int):
            self.rows[key] = Row.coerce(value, self.n_columns)
        elif isinstance(key, tuple) and len(key) == 2:
            row, column = key
            if isinstance(row, slice):
                value = tuple(value)
                start, stop, stride = row.indices(len(self.rows))
                for i, v in zip(range(start, stop, stride), value):
                    self[i, column] = v
            elif not isinstance(row, int):
                raise TypeError("Table index %r" % (key,))
            else:
                self.rows[row][column] = value
        else:
            raise IndexError("Table index %r" % (key,))

    # adding texstrs ---
    def cell(
            self,
            content: FMTextLike = None,
            tag: str = None,
            width: int = 1,
            just: str = None,
    ):
        """Add a cell to the table

        Parameters
        ----------
        content : FMText
            Cell content.
        tag : str
            Formatting tag.
        width : int
            Width in columns for multicolumn cells.
        just : 'l' | 'r' | 'c'
            Justification (default: use column standard).
        """
        if self._active_row is None or len(self._active_row) == self.n_columns:
            self._active_row = self.add_row()
        self._active_row.cell(content, tag, width, just)

    def cells(self, *cells, **kwargs):
        "Add several simple cells with one command"
        for cell in cells:
            self.cell(cell, **kwargs)

    def add_row(self, at=None):
        """Add a row without affecting the cursor

        Parameters
        ----------
        at : int
            Index at which to insert row (default is after the current row;
            midrules also count as rows).

        Returns
        -------
        row : Row
            the new row.
        """
        row = Row(self.n_columns)
        if at is None:
            self.rows.append(row)
        else:
            self.rows.insert(at, row)
        return row

    def empty_row(self):
        self.endline()
        self._active_row = self.add_row()

    def endline(self):
        "Finish the active row"
        if self._active_row is not None:
            for _ in range(self.n_columns - len(self._active_row)):
                self._active_row.append(Cell())
            self._active_row = None

    def midrule(self, span=None):
        """Add a midrule

        note that a toprule and a bottomrule are inserted automatically
        in every table.

        Parameters
        ----------
        span : str | sequence of int
            Add a midrule that spans less than the whole width of the table
            (e.g., ``'2-4'`` or ``(2, 4)``).
        """
        self.endline()
        if span is None:
            self.rows.append("\\midrule")
        else:
            if isinstance(span, (list, tuple)):
                span = '%i-%i' % span
            elif isinstance(span, str):
                if not re.match('\d+-\d+', span):
                    raise ValueError("span=%r" % span)
            else:
                raise TypeError("span=%r" % (span,))
            self.rows.append(r"\cmidrule{%s}" % span)

    def title(self, content: FMTextLike):
        """Set the table title"""
        self._title = asfmtext_or_none(content)

    def caption(self, content: FMTextLike):
        """Set the table caption"""
        self._caption = asfmtext_or_none(content)

    def __repr__(self):
        # return self.__str__ so that when a function returns a Table, the
        # result can be inspected without assigning the Table to a variable.
        return self.__str__()

    def get_html(self, env={}):
        if self._caption is None:
            caption = None
        else:
            if preferences['html_tables_in_fig']:
                tag = 'figcaption'
            else:
                tag = 'caption'
            caption = _html_element(tag, self._caption, env)

        # table body
        table = []
        if caption and not preferences['html_tables_in_fig']:
            table.append(caption)
        for row in self.rows:
            if isinstance(row, str):
                if row == "\\midrule":
                    pass
#                     table.append('<tr style="border-bottom:1px solid black">')
            else:
                table.append(row.get_html(env))
        body = '\n'.join(table)

        # table frame
        if self.rules:
            table_options = {'border': 1, 'frame': 'hsides', 'rules': 'none'}
        else:
            table_options = {'border': 0}
        table_options['cellpadding'] = 2
        txt = _html_element('table', body, env, table_options)

        # embedd in a figure
        if preferences['html_tables_in_fig']:
            if caption:
                txt = '\n'.join((txt, caption))
            txt = _html_element('figure', txt, env)

        return txt

    def get_rtf(self, env={}):
        # header
        rows = ['\cellx%i000' % i for i in range(len(self.columns))]
        rows.insert(0, '\\trowd')
        rows.append('\\row')
        # body
        for row in self.rows:
            if isinstance(row, str):
                if row == "\\midrule":
                    pass
            else:
                rows.append(row.get_rtf(env))
        return '\n'.join(rows)

    def get_str(self, env={}, delim='   ', linesep='\n'):
        """Convert Table to str

        Parameters
        ----------
        env : dict
            Processing environment.
        delim : str
            Delimiter between columns.
        linesep : str
            Line separation string
        """
        # append to recent tex out
        _add_to_recent(self)

        if len(self.rows) == 0:
            return ''

        # determine column widths
        widths = []
        for row in self.rows:
            if not isinstance(row, str):  # some commands are str
                row_strlen = row._strlen(env)
                while len(row_strlen) < self.n_columns:
                    row_strlen.append(0)
                widths.append(row_strlen)
        widths = np.array(widths, dtype=int)
        c_width = np.max(widths, axis=0)  # column widths!

        # FIXME: take into account tab length:
        midrule = delim.join(['-' * w for w in c_width])
        midrule = midrule.expandtabs(4).replace(' ', '-')

        # collect lines
        txtlines = []
        for row in self.rows:
            if isinstance(row, str):  # commands
                if row == r'\midrule':
                    txtlines.append(midrule)  # "_"*l_len)
                elif row == r'\bottomrule':
                    txtlines.append(midrule)  # "_"*l_len)
                elif row == r'\toprule':
                    txtlines.append(midrule)  # "_"*l_len)
                elif row.startswith(r'\cmidrule'):
                    txt = row.split('{')[1]
                    txt = txt.split('}')[0]
                    start, end = txt.split('-')
                    start = int(start) - 1
                    end = int(end)
                    line = [' ' * w for w in c_width[:start]]
                    rule = delim.join(['-' * w for w in c_width[start:end]])
                    rule = rule.expandtabs(4).replace(' ', '-')
                    line += [rule]
                    line += [' ' * w for w in c_width[start:end]]
                    txtlines.append(delim.join(line))
                else:
                    pass
            else:
                txtlines.append(row.get_str(c_width, self.columns, delim, env))
        out = txtlines

        if self._title is not None:
            out = ['', self._title.get_str(env), ''] + out

        if isinstance(self._caption, str):
            out.append(self._caption)
        elif self._caption:
            out.append(str(self._caption))

        return linesep.join(out)

    def get_tex(self, env=None):
        if env is None:
            return tex(self)
        standalone = env.get('standalone')  # https://stackoverflow.com/a/17235546
        # init
        items = []
        if not standalone:
            items.append(r"\begin{center}")
        items.append(r"\begin{tabular}{%s}" % self.columns)
        if self.rules:
            items.append(r"\toprule")
        # Body
        for row in self.rows:
            if isinstance(row, str):
                items.append(row)
            else:
                items.append(row.get_tex(env))
        # post
        if self.rules:
            items.append(r"\bottomrule")
        items.append(r"\end{tabular}")
        if not standalone:
            items.append(r"\end{center}")
        return '\n'.join(items)

    def get_tsv(self, delimiter='\t', linesep='\n', fmt='%.9g'):
        r"""
        Return the table as tab-separated values (TSV) string.

        Parameters
        ----------
        delimiter : str
            Delimiter between columns (default ``'\t'``).
        linesep : str
            Delimiter string between lines (default ``'\n'``).
        fmt : str
            Format string for numerical entries (default ``'%.9g'``).
        """
        table = []
        for row in self.rows:
            if isinstance(row, str):
                pass
            else:
                table.append(row.get_tsv(delimiter, fmt=fmt))
        return linesep.join(table)

    def save_docx(self, path=None):
        """Save table as *.docx (requires `python-docx <https://python-docx.readthedocs.io>`_)

        Parameters
        ----------
        path : str
            Target file name (leave unspecified to use Save As dialog).

        Notes
        -----
        Most style options are not implemented.
        """
        from docx import Document

        if path is None:
            path = ui.ask_saveas(
                "Save Text File", filetypes=[("Word Document (*.docx)", "*.docx")])
        if not path:
            return
        document = Document()
        if self._title:
            document.add_heading(str(self._title), 0)

        table = document.add_table(rows=0, cols=self.n_columns)

        # Body
        for row in self.rows:
            if isinstance(row, str):
                continue
            d_row = table.add_row()
            i = 0
            for cell in row:
                d_row.cells[i].text = str(cell)
                i += cell.width

        document.save(path)

    def save_tsv(self, path=None, delimiter='\t', linesep='\n', fmt='%.15g'):
        r"""
        Save the table as tab-separated values file.

        Parameters
        ----------
        path : str | None
            Destination file name.
        delimiter : str
            String that is placed between cells (default: ``'\t'``).
        linesep : str
            Delimiter string between lines (default ``'\n'``).
        fmt : str
            Format string for representing numerical cells.
            (see 'Python String Formatting Documentation
            <http://docs.python.org/library/stdtypes.html#string-formatting-operations>'_)
        """
        _save_txt(self.get_tsv(delimiter, linesep, fmt), path)

    def save_txt(self, path=None, fmt='%.15g', delim='   ', linesep='\n'):
        r"""
        Save the table as text file.

        Parameters
        ----------
        path : str | None
            Destination file name.
        fmt : str
            Format string for representing numerical cells (default '%.15g').
        delim : str
            Cell delimiter.
        linesep : str
            String that is placed in between lines (default is ``'\n'``).
        """
        _save_txt(self.get_str({'fmt': fmt}, delim, linesep), path)


class Image(FMTextElement, BytesIO):
    "Represent an image file"

    def __init__(self, name=None, format='png', alt=None, buf=b''):
        """Represent an image file

        Parameters
        ----------
        name : str
            Name for the file (without extension; default is 'image').
        format : str
            File format (default 'png').
        alt : None | str
            Alternate text, placeholder in case the image can not be found
            (HTML `alt` tag, default is ``name``).
        buf : bytes
            Image buffer (optional).
        """
        BytesIO.__init__(self, buf)

        self.name = name or 'image'
        self.format = format
        self._alt = alt or self.name
        self._filename = os.extsep.join((self.name, format))

    @classmethod
    def from_array(cls, array, name='array', format='png', alt=None):
        """Create an Image object from an array.

        Parameters
        ----------
        array : array_like
            RGBA image array.
        name : None | str
            Name for the target image.
        format : str
            Format to save (default ``'png'``).
        alt : None | str
            Alternate text, placeholder in case the image can not be found
            (HTML `alt` tag).
        """
        im = cls(name, format, alt)
        imsave(im, np.asarray(array), format=format)
        return im

    @classmethod
    def from_file(cls, path, name=None, alt=None):
        """Create an Image object from an existing image file.

        Parameters
        ----------
        path : str
            Path to the image file.
        name : str
            Name for the file (the default is ``os.path.basename(path)`` without
            extension.
        alt : None | str
            Alternate text, placeholder in case the image can not be found
            (HTML `alt` tag).
        """
        if name is None:
            filename = os.path.basename(path)
            name, ext = os.path.splitext(filename)
        else:
            _, ext = os.path.splitext(path)

        with open(path, 'rb') as fid:
            buf = fid.read()

        return cls(name, ext[1:], alt, buf)

    def _repr_items(self):
        out = [repr(self._filename)]
        if self._alt != self._filename:
            out.append(repr(self._alt))
        v = self.getvalue()
        if len(v) > 0:
            out.append('buf=%s...' % repr(v[:50]))
        return out

    def get_html(self, env={}):
        resource_dir = env.get('resource_dir', None)
        if resource_dir is None:
            buf = self.getvalue()
            if self.format == 'svg':
                return buf.decode()
            # http://stackoverflow.com/a/7389616/166700
            data = base64.b64encode(buf).decode().replace('\n', '')
            src = 'data:image/{};base64,{}'.format(self.format, data)
        else:
            dirpath = os.path.join(env['root'], resource_dir)
            abspath = os.path.join(dirpath, self._filename)
            if os.path.exists(abspath):
                i = 0
                name, ext = os.path.splitext(self._filename)
                while os.path.exists(abspath):
                    i += 1
                    filename = name + ' %i' % i + ext
                    abspath = os.path.join(dirpath, filename)

            self.save_image(abspath)
            src = os.path.relpath(abspath, env['root'])

        return '<img src="%s" alt="%s">' % (src, html(self._alt))

    def get_str(self, env={}):
        txt = "Image (%s)" % str(self._alt)
        return txt

    def save_image(self, dst):
        if os.path.isdir(dst):
            dst = os.path.join(dst, self._filename)

        buf = self.getvalue()
        with open(dst, 'wb') as fid:
            fid.write(buf)


class Figure(FMText):
    "Represent a figure"

    def __init__(
            self,
            content: FMTextLike,
            caption: FMTextLike = None,
            options: dict = None,
    ):
        """Represent a figure

        Parameters
        ----------
        content : FMTextLike
            Figure content.
        caption : FMTextLike
            Figure caption.
        options : dict
            HTML options for ``<figure>`` tag.
        """
        self._caption = asfmtext(caption)
        FMText.__init__(self, content, None, options)

    def get_html(self, env={}):
        body = FMText.get_html(self, env)
        if self._caption:
            caption = _html_element('figcaption', self._caption, env)
            body = '\n'.join((body, caption))
        return _html_element('figure', body, env, self.options)

    def get_str(self, env={}):
        body = FMText.get_str(self, env)
        if self._caption:
            caption = str(self._caption)
            return "%s\n\n%s\n" % (body, caption)
        else:
            return body + '\n'


class Section(FMText):

    def __init__(
            self,
            heading: FMTextLike,
            content: FMTextLike = None,
    ):
        """Represent a section of an FMText document

        Parameters
        ----------
        heading : FMTextLike
            Section heading.
        content : FMTextLike
            Section content. Can also be constructed dynamically through the
            different .add_... methods.
        """
        self._heading = asfmtext(heading)
        FMText.__init__(self, content)

    def add_figure(
            self,
            caption: FMTextLike,
            content: FMTextLike = None,
            options: dict = None,
    ):
        """Add a figure frame to the section

        Parameters
        ----------
        caption : FMTextLike
            Figure caption.
        content : FMTextLike
            Figure content.
        options : dict
            HTML options for ``<figure>`` tag.

        Returns
        -------
        figure : Figure
            Figure object that was added.

        See Also
        --------
        .add_image_figure: add figure with a single image
        """
        if content is None:
            content = []
        figure = Figure(content, caption, options)
        self.append(figure)
        return figure

    def add_image_figure(self, image, caption, alt=None):
        """Add an image in a figure frame to the section

        Parameters
        ----------
        image : Image | array | str
            Image, image array or target filename for the image. If a filename
            it should have the appropriate extension.
        caption : FMText
            Figure caption.
        alt : None | str
            Alternate text, placeholder in case the image can not be found
            (HTML `alt` tag).

        Returns
        -------
        image : Image
            Image object that was added.

        See Also
        --------
        .add_figure: add an empty figure

        Notes
        -----
        Return the Image object which can be used to write the image data. If
        a function supports writing to a file-like object, it can be written
        with ``save_image(image)``. A function that writes a file to disk can
        be used as ``save_image(image.``.

        """
        if isinstance(image, str):
            name, ext = os.path.splitext(image)
            if ext:
                format = ext[1:]
            else:
                format = 'png'
            image = Image(name, format, alt)
        elif isinstance(image, np.ndarray):
            image = Image.from_array(image, alt=alt)
        else:
            image = asfmtext(image)

        figure = Figure(image, caption)
        self.append(figure)
        return image

    def add_paragraph(self, content=[]):
        paragraph = FMText(content, 'paragraph')
        self.append(paragraph)
        return paragraph

    def add_section(
            self,
            heading: FMTextLike,
            content: FMTextLike = ()):
        """Add a new subordinate section

        Parameters
        ----------
        heading : FMTextLike
            Heading for the section.
        content : FMTextLike
            Content for the section.

        Returns
        -------
        section : Section
            The new section.
        """
        section = Section(heading, content)
        self.append(section)
        return section

    def get_html(self, env={}):
        env = env.copy()

        heading = self._get_html_section_heading(env)
        body = FMText.get_html(self, env)
        return '\n\n'.join(('', heading, body))

    def _get_html_section_heading(self, env):
        heading = self._heading.get_html(env)

        level = env.get('level', 1)
        tag = 'h%i' % level
        txt = _html_element(tag, heading, env)
        if 'toc' in env:
            toc_id = max(env['toc_ids']) + 1
            env['toc_ids'].append(toc_id)
            txt = _html_element('a', txt, env, {'name': toc_id})

            toc_txt = _html_element('a', heading, env, {'href': '#%i' % toc_id})
            env['toc'].append((level, toc_txt))
        env['level'] = level + 1
        return txt

    def get_str(self, env={}):
        level = env.get('level', (1,))
        number = '.'.join(map(str, level))
        title = ' '.join((number, str(self._heading)))
        if len(level) == 1:
            underline_char = '='
        else:
            underline_char = '-'
        underline = underline_char * len(title)

        content = [title, underline, '']
        env = env.copy()
        level = list(level) + [1]
        for item in self.content:
            if isinstance(item, Section):
                env['level'] = tuple(level)
                txt = item.get_str(env)
                level[-1] += 1
                content += ['', '', txt]
            else:
                content += [str(item)]

        txt = '\n'.join(content)
        return txt


class Report(Section):

    def __init__(
            self,
            title: FMTextLike,
            author: FMTextLike = None,
            date: Union[FMTextLike, bool] = True,
            content: FMTextLike = None,
            site_title: str = None,
    ):
        """Represent an FMText report document

        Parameters
        ----------
        title : FMTextLike
            Document title.
        author : FMTextLike
            Document autho.
        date : bool | FMTextLike
            Date to print on the report. If True (default), the current day
            (object initialization) is used.
        content : FMTextLike
            Report content. Can also be constructed dynamically through the
            different .add_... methods.
        site_title : str
            Set the HTML site title (the default is the same as title).
        """
        if author is not None:
            author = asfmtext(author, r'\author')
        if date:
            if date is True:
                date = str(datetime.date.today())
            elif isinstance(date, str) and '%' in date:
                date = datetime.datetime.now().strftime(date)
            date = asfmtext(date, r'\date')
        self._author = author
        self._date = date
        self._site_title = site_title
        Section.__init__(self, title, content)

    def _repr_items(self):
        out = list(map(repr, (self._heading, self._author, self._date, self.content)))
        if self._site_title:
            out.append(repr(self._site_title))
        return out

    def get_html(self, env={}):
        # setup TOC in env
        env = env.copy()
        env['toc'] = []
        env['toc_ids'] = [-1]
        env['level'] = 2

        # format document body (& collect document info)
        body = FMText.get_html(self, env)

        # format TOC
        toc = ['<ul>']
        level = 2
        for item_level, item in env['toc']:
            if item_level > level:
                toc.append('<ul>' * (item_level - level))
            elif item_level < level:
                toc.append('</ul>' * (level - item_level))
            toc.append(_html_element('li', item, env))
            level = item_level
        toc.append('</ul></li>' * (level - 1) + '</ul>')
        toc = '\n'.join(toc)

        # compile document content
        content = []
        title = _html_element('h1', self._heading, env)
        content.append(title)
        if self._author is not None:
            author = html(self._author, env)
            content.append(author)
        if self._date:
            date = html(self._date, env)
            content.append(date)
        content.append(toc)
        content.append(body)
        txt = '\n<br>\n'.join(content)
        return txt

    def get_str(self, env={}):
        content = []
        title = str(self._heading)
        underline = '^' * len(title)
        content += [title, underline, '']
        if self._author is not None:
            author = self._author.get_str(env)
            content += [author, '']
        if self._date:
            date = self._date.get_str(env)
            content += [date, '']

        level = [1]
        env = env.copy()
        for item in self.content:
            if isinstance(item, Section):
                env['level'] = tuple(level)
                txt = item.get_str(env)
                level[-1] += 1
            else:
                txt = item.get_str(env)
            content += [txt, '']

        txt = '\n'.join(content)
        return txt

    def pickle(self, path, extension='.pickled'):
        """Pickle the Report object

        Parameters
        ----------
        path : str
            Location where to save the report. For None, the file is saved
            in the report's folder.
        extension : None | str
            Extension to append to the path. If extension is None, or path
            already ends with extension nothing is done.
        """
        if extension and not path.endswith(extension):
            path += extension

        with open(path, 'wb') as fid:
            pickle.dump(self, fid, pickle.HIGHEST_PROTOCOL)

    def save_html(self, path, embed_images=True, meta=None):
        """Save HTML file of the report

        Parameters
        ----------
        path : str
            Path at which to save the html. Does not need to contain the
            extension. A folder with the same name is created fo resource
            files.
        embed_images : bool
            Embed images in the HTML file (default True). If False, a separate
            folder containing image files is created.
        meta : dict
            Meta-information for document head.
        """
        if path.endswith('.html'):
            path = path[:-5]

        save_html(self, path, embed_images, meta)

    def show(self):
        "Save the report as temporary file and open in browser"
        temp_dir = tempfile.mkdtemp()
        path = os.path.join(temp_dir, 'report.html')
        self.save_html(path)
        webbrowser.open('file:/' + os.path.realpath(path))

    def sign(self, packages=('eelbrain',)):
        """Add a signature to the report

        Parameters
        ----------
        packages : iterator of str
            Packages whose ``__version__`` to include in the signature (default
            is eelbrain only).

        Notes
        -----
        Includes informationon on the computer that created the report, time
        of creation, and version of selected packages.
        """
        info = ["Created by %s on %s" % (socket.gethostname(),
                                         time.strftime("%c"))]
        for package in packages:
            if package == 'eelbrain':
                name = 'Eelbrain'
            elif package == 'mne':
                name = 'MNE-Python'
            elif package == 'surfer':
                name = 'PySurfer'
            else:
                name = package

            mod = import_module(package)
            text = "%s version %s" % (name, mod.__version__)
            info.append(text)

        signature = ' \u2014 \n'.join(info)
        self.add_paragraph(signature)


def symbol(symbol, subscript=None, tag='math'):
    if subscript is None:
        return Text(symbol, tag)
    else:
        return FMText([Text(symbol), Text(subscript, '_')], tag)


def p(p, stars=None, of=3, tag='math'):
    """:class:`FMText` representation of a p-value

    Parameters
    ----------
    p : scalar
        P-value.
    stars : None | str
        Stars decorating the p-value (e.g., "**")
    of : int
        Max numbers of star characters possible (to add empty space for
        alignment).

    Returns
    -------
    text : FMText
        FMText with formatted p-value.
    """
    if stars is None:
        return P(p)
    else:
        return FMText([P(p), Stars(stars, of=of)], tag)


def stat(x, fmt="%.2f", stars=None, of=3, tag='math', drop0=False):
    ":class:`FMText` with properties for a statistic (e.g. a t-value)"
    if stars is None:
        return Number(x, tag, fmt, drop0)
    elif isinstance(stars, float):
        stars_obj = Stars.from_p(stars, of)
    else:
        stars_obj = Stars(stars, of)
    return FMText([Number(x, None, fmt, drop0), stars_obj], tag)


def eq(name, result, subscript=None, fmt='%.2f', stars=None, of=3, drop0=False):
    "``$name_{subscript} = result^stars$``"
    symbol_ = symbol(name, subscript, None)
    stat_ = stat(result, fmt, stars, of, None, drop0)
    return FMText([symbol_, Text(' = '), stat_], 'math')


def peq(p, subscript=None, stars=None, of=3):
    symbol_ = symbol('p', subscript, None)

    if p < .001:
        eq_ = Text(' ')
    else:
        eq_ = Text(' = ')

    if stars is None:
        return FMText([symbol_, eq_, P(p)], 'math')
    else:
        if stars is True:
            stars_obj = Stars.from_p(p, of)
        else:
            stars_obj = Stars(stars, of)
        return FMText([symbol_, eq_, P(p), stars_obj], 'math')


def delim_list(items, delimiter=', '):
    delim = asfmtext(delimiter)
    out = list(items)
    for i in range(len(out) - 1, 0, -1):
        out.insert(i, delim)
    return out


def ms(t_s):
    "Convert time in seconds to rounded milliseconds"
    return int(round(t_s * 1000))


def ms_window(t0, t1):
    return f"{ms(t0)} - {ms(t1)} ms"


def unindent(text, skip1=False):
    """Remove leading spaces that are present in all lines of ``text``.

    Parameters
    ----------
    test : str
        The text from which leading spaces should be removed.
    skip1 : bool
        Ignore the first line when determining number of spaces to unindent,
        and remove all leading whitespaces from it.
    """
    # count leading whitespaces
    lines = text.splitlines()
    ws_lead = []
    for line in lines[skip1:]:
        len_stripped = len(line.lstrip(' '))
        if len_stripped:
            ws_lead.append(len(line) - len_stripped)

    if len(ws_lead) > skip1:
        rm = min(ws_lead)
        if rm:
            if skip1:
                lines[0] = ' ' * rm + lines[0].lstrip()

            text = '\n'.join(line[rm:] for line in lines)

    return text


def im_table(ims, header=None, name="im_table"):
    """Create an SVG with a table of images

    Parameters
    ----------
    ims : list of list of arrays
        Images. Should all have same shape.
    header : None | list of str
        List of column titles.

    Returns
    -------
    images : Image
        FMTXT Image object that can be saved as SVG or integrated into an
        FMTXT document.
    """
    lens = set(map(len, ims))
    if len(lens) > 1:
        raise ValueError("Unequal number of columns")
    n_cols = lens.pop()
    n_rows = len(ims)

    # test im shape
    im0 = ims[0][0]
    if im0.ndim != 3:
        raise ValueError("Images need ndim=3, got %i" % ims[0][0].ndim)
    shape = im0.shape
    if not all(im.shape == shape for line in ims for im in line):
        raise NotImplementedError("Not all images have same shape")
    im_h, im_w, _ = shape

    xs = list(range(0, im_w * n_cols, im_w))
    y0 = 0 if header is None else 25
    ys = list(range(y0, y0 + im_h * n_rows, im_h))

    svg_h = y0 + im_h * n_rows
    svg_w = im_w * n_cols
    svg = ['<svg width="{w}" height="{h}">'.format(w=svg_w, h=svg_h)]
    if header is not None:
        assert len(header) == n_cols
        p = '<text x="{x}" y="{y}">{txt}</text>'
        for x, txt in zip(xs, header):
            item = p.format(x=x, y=18, txt=txt)
            svg.append(item)

    p = ('<image x="{x}" y="{y}" width="{w}" height="{h}" xlink:href='
         '"data:image/png;base64,{data}" />')
    for y, line in zip(ys, ims):
        for x, im in zip(xs, line):
            data = _array_as_png(im)
            item = p.format(x=x, y=y, w=im_w, h=im_h, data=data)
            svg.append(item)

    svg.append("</svg>")
    buf = '\n'.join(svg)
    return Image(name, 'svg', buf=buf.encode())


def _array_as_png(im):
    """Convert array to base64 encoded PNG

    Parameters
    ----------
    im : array_like
        An MxN (luminance), MxNx3 (RGB) or MxNx4 (RGBA) array.

    Returns
    -------
    data : str
        Encoded PNG image.
    """
    buf = BytesIO()
    imsave(buf, np.asarray(im), format='png')
    data = buf.getvalue()
    return data


class MetaParser(HTMLParser):
    def __init__(self, file_name):
        HTMLParser.__init__(self)
        self.out = None
        with open(file_name) as fid:
            while self.out is None:
                self.feed(fid.readline())

    def handle_starttag(self, tag, attrs):
        if tag == 'meta':
            self.out = dict(attrs)

    def handle_endtag(self, tag):
        "No need to keep reading"
        if tag == 'head':
            self.out = {}


def read_meta(file_path):
    "Read meta information from a HTML head"
    return MetaParser(file_path).out
